#!/usr/bin/env python3
import argparse
import html
import json
import shutil
import sys
import subprocess
import zipfile
import os
from pathlib import Path


def clean_visible_text(value):
    text = str(value or '')
    for prefix in ('匹配 JD：', '匹配JD：', '匹配：'):
        if text.startswith(prefix):
            text = text[len(prefix):].strip()
    replacements = {
        '匹配策略': '培养方向',
        '投递关键词': '核心能力',
        'JD 反推': '岗位理解',
        'ATS 投喂': '筛选友好表达',
        '缺口分析': ''
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text.strip()


def esc(value):
    return html.escape(clean_visible_text(value), quote=True)


def safe_name(value):
    keep = []
    for ch in str(value or ''):
        if ch.isalnum() or ch in ('-', '_') or '一' <= ch <= '龥':
            keep.append(ch)
        elif ch.isspace():
            keep.append('-')
    name = ''.join(keep).strip('-')
    return name or '毕业生岗位定制简历'


def get_design(data):
    target = data.get('target', {})
    brief = data.get('design_brief') or {}
    style = target.get('style') or ''
    default_palette = {
        'primary': '#1f3f68',
        'accent': '#2b87c8',
        'soft': '#eef6fb',
        'paper': '#ffffff',
        'ink': '#1d2730',
        'muted': '#64717d',
        'line': '#d9e3ea',
        'background': '#edf2f5'
    }
    palette = {**default_palette, **(brief.get('palette') or {})}
    layout = brief.get('layout') or style or 'top-card'
    if layout == 'professional':
        layout = 'top-card'
    if layout == 'creative':
        layout = 'portfolio-light'
    return {
        'layout': layout,
        'tone': brief.get('tone') or '专业、清晰、可投递',
        'palette': palette,
        'density': brief.get('density') or 'balanced',
        'print_scale': float(brief.get('print_scale') or 1),
        'photo': brief.get('photo') or {},
        'module_order': brief.get('module_order') or [],
        'avoid': brief.get('avoid') or []
    }


def contact_items(data):
    basics = data.get('basics', {})
    items = []
    for label, value in [('电话', basics.get('phone')), ('邮箱', basics.get('email')), ('城市', basics.get('location'))]:
        if value:
            items.append((label, value))
    for link in basics.get('links') or []:
        items.append(('链接', link))
    return items


def infer_profile(data):
    basics = data.get('basics', {})
    target = data.get('target', {})
    rows = []
    for label, value in [
        ('目标岗位', target.get('role') or basics.get('title')),
        ('目标公司', target.get('company')),
        ('行业方向', target.get('industry')),
        ('所在城市', basics.get('location'))
    ]:
        if value:
            rows.append((label, value))
    return rows


def render_photo(data, design):
    basics = data.get('basics', {})
    photo = basics.get('photo')
    photo_cfg = design.get('photo') or {}
    if not photo or photo_cfg.get('mode') in {'none', 'ask'}:
        return ''
    return f'<div class="photo"><img src="{esc(photo)}" alt="简历照片"></div>'


def render_chips(items):
    return ''.join(f'<span class="chip">{esc(x)}</span>' for x in (items or []) if x)


def render_summary_cards(summary):
    labels = ['岗位意愿', '能力基础', '发展方向']
    cards = []
    for i, text in enumerate(summary[:3]):
        label = labels[i] if i < len(labels) else '核心优势'
        cards.append(f'<div class="summary-card"><b>{esc(label)}</b><p>{esc(text)}</p></div>')
    return ''.join(cards)


def render_sections(sections, timeline=False):
    out = []
    for section in sections:
        items = []
        for item in section.get('items', []):
            bullets = ''.join(f'<li>{esc(b)}</li>' for b in item.get('bullets', []) if b)
            sub = item.get('subheading')
            meta = item.get('meta')
            item_class = 'item timeline-item' if timeline else 'item'
            items.append(f'''<article class="{item_class}">
  <div class="item-head"><h3>{esc(item.get('heading'))}</h3><span>{esc(meta)}</span></div>
  {f'<p class="subheading">{esc(sub)}</p>' if sub else ''}
  {f'<ul>{bullets}</ul>' if bullets else ''}
</article>''')
        out.append(f'<section><h2>{esc(section.get("title"))}</h2>{"".join(items)}</section>')
    return ''.join(out)


def css_tokens(design):
    p = design['palette']
    return f''':root {{
  --primary: {p.get('primary')};
  --accent: {p.get('accent')};
  --soft: {p.get('soft')};
  --paper: {p.get('paper')};
  --ink: {p.get('ink')};
  --muted: {p.get('muted')};
  --line: {p.get('line')};
  --bg: {p.get('background')};
}}'''


def base_css(design):
    layout = design['layout']
    print_scale = max(0.72, min(1.0, float(design.get('print_scale') or 1)))
    if layout == 'sidebar-card':
        layout_css = '''
.page { overflow: hidden; }
.hero { background: linear-gradient(135deg, var(--primary), color-mix(in srgb, var(--primary) 76%, #ffffff)); color: white; padding: 17mm 18mm 13mm; display: grid; grid-template-columns: 1.15fr .85fr; gap: 18px; align-items: end; }
.title-line { color: color-mix(in srgb, var(--accent) 65%, #ffffff); }
.hero-note { color: rgba(255,255,255,.84); }
.contact-card { justify-self: end; width: 212px; border: 1px solid rgba(255,255,255,.32); background: rgba(255,255,255,.10); padding: 13px 14px; }
.contact-card p, .contact-card b { color: white; }
.content { display: grid; grid-template-columns: 58mm 1fr; min-height: calc(297mm - 47mm); }
aside { background: var(--soft); padding: 11mm 8mm 14mm 12mm; border-right: 1px solid var(--line); }
main { padding: 11mm 14mm 14mm; }
.summary-grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 9px; }
.summary-card { border: 1px solid var(--line); border-top: 3px solid var(--accent); background: #fff; padding: 9px; min-height: 82px; }
'''
    elif layout == 'timeline':
        layout_css = '''
.hero { border-top: 9px solid var(--primary); padding: 16mm 18mm 10mm; background: var(--paper); }
.content { padding: 0 18mm 16mm; }
main { display: block; }
.timeline-item { border-left: 2px solid var(--line); padding-left: 14px; }
.timeline-item::before { width: 9px; height: 9px; border-radius: 50%; left: -6px; top: 7px; }
'''
    elif layout == 'portfolio-light':
        layout_css = '''
.hero { padding: 16mm 18mm 11mm; background: linear-gradient(120deg, var(--soft), #fff); border-bottom: 4px solid var(--accent); display: grid; grid-template-columns: 1fr auto; gap: 18px; }
.content { padding: 12mm 18mm 16mm; }
main { display: block; }
.summary-grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 10px; }
.summary-card { background: var(--soft); border: 1px solid var(--line); padding: 10px; }
'''
    elif layout == 'technical-grid':
        layout_css = '''
.hero { padding: 15mm 18mm 10mm; background: #111827; color: #fff; }
.title-line { color: color-mix(in srgb, var(--accent) 72%, #ffffff); }
.hero-note { color: rgba(255,255,255,.78); }
.content { padding: 11mm 18mm 16mm; }
main { display: block; }
.summary-grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 9px; }
.summary-card { background: var(--soft); border-left: 4px solid var(--accent); padding: 10px; }
'''
    else:
        layout_css = '''
.hero { padding: 15mm 18mm 10mm; background: var(--paper); border-top: 7px solid var(--primary); border-bottom: 1px solid var(--line); }
.content { padding: 11mm 18mm 16mm; }
main { display: block; }
.summary-grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 9px; }
.summary-card { border: 1px solid var(--line); background: var(--soft); padding: 10px; }
'''
    return f'''{css_tokens(design)}
* {{ box-sizing: border-box; }}
body {{ margin: 0; background: var(--bg); color: var(--ink); font-family: "PingFang SC", "Microsoft YaHei", "Noto Sans CJK SC", Arial, sans-serif; line-height: 1.48; }}
.page {{ width: 210mm; min-height: 297mm; margin: 18px auto; background: var(--paper); box-shadow: 0 24px 70px rgba(31, 41, 55, .16); }}
h1 {{ margin: 0; font-size: 37px; line-height: 1; letter-spacing: 0; }}
.title-line {{ margin-top: 9px; font-size: 16px; font-weight: 800; color: var(--primary); }}
.hero-note {{ margin-top: 10px; max-width: 560px; color: var(--muted); font-size: 12.4px; }}
.contact-card b {{ display: block; margin-bottom: 7px; font-size: 12px; color: var(--primary); }}
.contact-card p {{ margin: 4px 0; font-size: 12px; color: var(--ink); }}
.photo {{ width: 78px; height: 96px; border: 3px solid #fff; box-shadow: 0 8px 24px rgba(0,0,0,.16); overflow: hidden; background: var(--soft); }}
.photo img {{ width: 100%; height: 100%; object-fit: cover; display: block; }}
.side-block {{ margin-bottom: 17px; }}
.side-title, h2 {{ margin: 0 0 9px; font-size: 13px; color: var(--primary); display: flex; align-items: center; gap: 8px; }}
.side-title::before, h2::before {{ content: ""; width: 6px; height: 17px; background: var(--accent); display: inline-block; }}
.profile-line {{ padding: 7px 0; border-bottom: 1px solid var(--line); }}
.profile-line b {{ display: block; font-size: 11px; color: var(--muted); font-weight: 500; }}
.profile-line span {{ display: block; margin-top: 2px; font-size: 12.3px; font-weight: 700; color: var(--ink); }}
.chips {{ display: flex; flex-wrap: wrap; gap: 7px; }}
.chip {{ font-size: 11.1px; color: var(--primary); border: 1px solid color-mix(in srgb, var(--accent) 36%, #ffffff); background: #fff; padding: 4px 7px; border-radius: 2px; }}
section {{ margin-bottom: 16px; }}
.summary-card b {{ display: block; font-size: 12.2px; color: var(--primary); margin-bottom: 5px; }}
.summary-card p {{ margin: 0; color: #44505a; font-size: 11.2px; }}
.item {{ position: relative; padding-left: 16px; margin: 11px 0 0; break-inside: avoid; }}
.item::before {{ content: ""; position: absolute; left: 1px; top: 6px; width: 8px; height: 8px; background: var(--accent); }}
.item-head {{ display: grid; grid-template-columns: 1fr auto; gap: 12px; align-items: baseline; }}
h3 {{ margin: 0; font-size: 13.6px; color: var(--ink); }}
.item-head span, .subheading {{ color: var(--muted); font-size: 11.3px; margin: 2px 0 0; }}
ul {{ margin: 5px 0 0; padding-left: 16px; }}
li {{ margin: 3px 0; font-size: 11.9px; }}
@page {{ size: A4; margin: 0; }}
@media print {{ body {{ background: #fff; }} .page {{ margin: 0; box-shadow: none; width: 210mm; min-height: auto; zoom: {print_scale}; }} .content {{ min-height: auto !important; }} }}
{layout_css}'''


def render_profile_panel(data):
    rows = infer_profile(data)
    return '<div class="side-block"><h2 class="side-title">个人信息</h2>' + ''.join(
        f'<div class="profile-line"><b>{esc(label)}</b><span>{esc(value)}</span></div>' for label, value in rows
    ) + '</div>'


def render_contact_card(data, design):
    photo = render_photo(data, design)
    contacts = ''.join(f'<p>{esc(label)}：{esc(value)}</p>' for label, value in contact_items(data))
    if photo:
        return f'<div style="display:flex; gap:12px; align-items:flex-start; justify-content:flex-end;">{photo}<div class="contact-card"><b>联系方式</b>{contacts}</div></div>'
    return f'<div class="contact-card"><b>联系方式</b>{contacts}</div>'


def render_html(data):
    basics = data.get('basics', {})
    target = data.get('target', {})
    design = get_design(data)
    name = basics.get('name') or '姓名'
    role = target.get('role') or basics.get('title') or '目标岗位'
    company = target.get('company')
    title = f'{role}' + (f'｜{company}' if company else '')
    summary = data.get('summary') or []
    sections = data.get('sections') or []
    keywords = data.get('keywords') or []
    layout = design['layout']
    hero_note = summary[0] if summary else design.get('tone', '')

    if layout == 'sidebar-card':
        hero = f'''<header class="hero">
  <div><h1>{esc(name)}</h1><div class="title-line">应聘岗位：{esc(title)}</div><div class="hero-note">{esc(hero_note)}</div></div>
  {render_contact_card(data, design)}
</header>'''
        aside = render_profile_panel(data)
        if keywords:
            aside += f'<div class="side-block"><h2 class="side-title">能力标签</h2><div class="chips">{render_chips(keywords[:8])}</div></div>'
        summary_html = f'<section><h2>岗位摘要</h2><div class="summary-grid">{render_summary_cards(summary)}</div></section>' if summary else ''
        chips_section = f'<section><h2>核心能力</h2><div class="chips">{render_chips(keywords)}</div></section>' if keywords else ''
        body = f'<div class="content"><aside>{aside}</aside><main>{summary_html}{render_sections(sections)}{chips_section}</main></div>'
    else:
        contacts = contact_items(data)
        contact_line = '<div class="chips" style="margin-top:10px;">' + ''.join(f'<span class="chip">{esc(label)}：{esc(value)}</span>' for label, value in contacts) + '</div>' if contacts else ''
        hero = f'''<header class="hero">
  <div><h1>{esc(name)}</h1><div class="title-line">应聘岗位：{esc(title)}</div><div class="hero-note">{esc(hero_note)}</div>{contact_line}</div>
  {render_photo(data, design)}
</header>'''
        summary_html = f'<section><h2>岗位摘要</h2><div class="summary-grid">{render_summary_cards(summary)}</div></section>' if summary else ''
        chips_section = f'<section><h2>核心能力</h2><div class="chips">{render_chips(keywords)}</div></section>' if keywords else ''
        body = f'<div class="content"><main>{summary_html}{render_sections(sections, timeline=(layout == "timeline"))}{chips_section}</main></div>'

    return f'''<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{esc(name)} - {esc(role)} - 简历</title>
<style>{base_css(design)}</style>
</head>
<body>
<div class="page">{hero}{body}</div>
</body>
</html>'''


def docx_escape(text):
    return esc(text).replace('\n', ' ')


def paragraph(text, style=None):
    style_xml = f'<w:pStyle w:val="{style}"/>' if style else ''
    return f'<w:p><w:pPr>{style_xml}</w:pPr><w:r><w:t xml:space="preserve">{docx_escape(text)}</w:t></w:r></w:p>'


def build_document_xml(data):
    basics = data.get('basics', {})
    target = data.get('target', {})
    parts = [paragraph(basics.get('name') or '姓名', 'Title')]
    contacts = [basics.get('title') or target.get('role'), basics.get('phone'), basics.get('email'), basics.get('location')]
    contacts += basics.get('links') or []
    parts.append(paragraph(' · '.join(str(x) for x in contacts if x)))
    if data.get('summary'):
        parts.append(paragraph('核心优势', 'Heading1'))
        for line in data['summary']:
            parts.append(paragraph('• ' + str(line)))
    for section in data.get('sections', []):
        parts.append(paragraph(section.get('title', ''), 'Heading1'))
        for item in section.get('items', []):
            head = item.get('heading', '')
            meta = item.get('meta')
            parts.append(paragraph(head + (f'｜{meta}' if meta else ''), 'Heading2'))
            if item.get('subheading'):
                parts.append(paragraph(item['subheading']))
            for bullet in item.get('bullets', []):
                parts.append(paragraph('• ' + str(bullet)))
    body = ''.join(parts)
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>{body}<w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="900" w:right="850" w:bottom="900" w:left="850"/></w:sectPr></w:body>
</w:document>'''


def write_docx_minimal(data, path):
    content_types = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>'''
    rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>'''
    styles = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:rPr><w:b/><w:sz w:val="34"/></w:rPr><w:pPr><w:spacing w:after="120"/></w:pPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="Heading 1"/><w:rPr><w:b/><w:color w:val="2454A6"/><w:sz w:val="24"/></w:rPr><w:pPr><w:spacing w:before="180" w:after="80"/></w:pPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="Heading 2"/><w:rPr><w:b/><w:sz w:val="22"/></w:rPr><w:pPr><w:spacing w:before="80" w:after="40"/></w:pPr></w:style>
</w:styles>'''
    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('[Content_Types].xml', content_types)
        zf.writestr('_rels/.rels', rels)
        zf.writestr('word/document.xml', build_document_xml(data))
        zf.writestr('word/styles.xml', styles)



def write_docx(data, path):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except Exception:
        write_docx_minimal(data, path)
        return

    # Word is an editable document, not the visual master. Keep it conservative so
    # it survives WPS, Word, Pages, and recruiter-side edits.
    basics = data.get('basics', {})
    target = data.get('target', {})
    design = get_design(data)
    palette = design['palette']

    def hex_clean(value, default):
        return str(value or default).replace('#', '').upper()

    primary = hex_clean(palette.get('primary'), '8F1F17')
    accent = hex_clean(palette.get('accent'), 'F4BD2A')
    ink = hex_clean(palette.get('ink'), '1D2730')
    muted = hex_clean(palette.get('muted'), '64717D')

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    for style_name in ['Normal', 'List Bullet']:
        style = doc.styles[style_name]
        style.font.name = 'Arial Unicode MS'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
        style.font.size = Pt(9)

    def add_run(paragraph, text, size=9, color=ink, bold=False):
        run = paragraph.add_run(clean_visible_text(text))
        run.font.name = 'Arial Unicode MS'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(hex_clean(color, ink))
        run.bold = bold
        return run

    def set_bottom_border(paragraph, color='D9C8B8', size='8'):
        p_pr = paragraph._p.get_or_add_pPr()
        borders = p_pr.find(qn('w:pBdr'))
        if borders is None:
            borders = OxmlElement('w:pBdr')
            p_pr.append(borders)
        bottom = borders.find(qn('w:bottom'))
        if bottom is None:
            bottom = OxmlElement('w:bottom')
            borders.append(bottom)
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), size)
        bottom.set(qn('w:space'), '3')
        bottom.set(qn('w:color'), color)

    name = basics.get('name') or '姓名'
    role = target.get('role') or basics.get('title') or '目标岗位'
    company = target.get('company')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, name, size=22, color=primary, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    title = f'应聘岗位：{role}' + (f'｜{company}' if company else '')
    add_run(p, title, size=11, color=primary, bold=True)

    if target.get('industry'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f"行业方向：{target['industry']}", size=8.5, color=muted)

    contact_values = []
    for label, value in contact_items(data):
        contact_values.append(f'{label}：{value}')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_bottom_border(p, color=accent, size='10')
    add_run(p, '联系方式：' + '  ·  '.join(contact_values), size=8.5, color=muted)

    def heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(3)
        set_bottom_border(p, color='EAD8C9', size='6')
        add_run(p, '▌ ', size=11, color=accent, bold=True)
        add_run(p, text, size=11, color=primary, bold=True)

    if data.get('summary'):
        heading('核心优势')
        for line in data['summary']:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.48)
            p.paragraph_format.first_line_indent = Cm(-0.18)
            p.paragraph_format.space_after = Pt(1)
            add_run(p, line, size=8.8, color=ink)

    for section_data in data.get('sections', []):
        heading(section_data.get('title') or '')
        for item in section_data.get('items', []):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(1)
            add_run(p, item.get('heading') or '', size=9.4, color=ink, bold=True)
            if item.get('meta'):
                add_run(p, '    ' + clean_visible_text(item['meta']).replace('->', '→'), size=8, color=muted)
            for bullet in item.get('bullets') or []:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Cm(0.48)
                p.paragraph_format.first_line_indent = Cm(-0.18)
                p.paragraph_format.space_after = Pt(0.8)
                p.paragraph_format.line_spacing = 1.03
                add_run(p, bullet, size=8.35, color=ink)

    if data.get('keywords'):
        heading('核心能力')
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_run(p, ' / '.join(clean_visible_text(kw) for kw in data['keywords']), size=8.5, color=primary, bold=True)

    doc.save(path)


def write_pdf_reportlab(data, pdf_path):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen import canvas
        from reportlab.platypus import Paragraph
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_LEFT
    except Exception:
        return False

    font_candidates = [
        '/System/Library/Fonts/Supplemental/Arial Unicode.ttf',
        '/Library/Fonts/Arial Unicode.ttf',
        '/System/Library/Fonts/Supplemental/Songti.ttc',
        '/System/Library/Fonts/STHeiti Light.ttc',
    ]
    font_path = next((p for p in font_candidates if Path(p).exists()), None)
    if not font_path:
        return False

    try:
        pdfmetrics.registerFont(TTFont('ResumeCN', font_path))
    except Exception:
        return False

    def color(value, default):
        return colors.HexColor(str(value or default))

    def draw_paragraph(c, text, x, y, width, style):
        p = Paragraph(clean_visible_text(text), style)
        _, h = p.wrap(width, 200)
        p.drawOn(c, x, y - h)
        return h

    basics = data.get('basics', {})
    target = data.get('target', {})
    summary = data.get('summary') or []
    sections = data.get('sections') or []
    keywords = data.get('keywords') or []
    design = get_design(data)
    palette = design['palette']

    W, H = A4
    mm = 72 / 25.4
    primary = color(palette.get('primary'), '#8f1f17')
    accent = color(palette.get('accent'), '#f4bd2a')
    soft = color(palette.get('soft'), '#fff2df')
    paper = color(palette.get('paper'), '#fffaf2')
    line = color(palette.get('line'), '#ead8c9')
    ink = color(palette.get('ink'), '#1d2730')
    muted = color(palette.get('muted'), '#64717d')
    white = colors.white
    font = 'ResumeCN'

    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    c.setTitle(f"{basics.get('name') or '姓名'} - {target.get('role') or basics.get('title') or '简历'}")
    c.setFillColor(paper)
    c.rect(0, 0, W, H, fill=1, stroke=0)

    hero_h = 48 * mm
    c.setFillColor(primary)
    c.rect(0, H - hero_h, W, hero_h, fill=1, stroke=0)
    c.setFillColor(white)
    c.setFont(font, 26)
    c.drawString(18 * mm, H - 22 * mm, clean_visible_text(basics.get('name') or '姓名'))
    title = target.get('role') or basics.get('title') or '目标岗位'
    company = target.get('company')
    if company:
        title = f'{title}｜{company}'
    c.setFont(font, 12)
    c.setFillColor(accent)
    c.drawString(18 * mm, H - 32 * mm, f'应聘岗位：{clean_visible_text(title)}')
    style_hero = ParagraphStyle('hero', fontName=font, fontSize=8.5, leading=12, textColor=white, alignment=TA_LEFT, wordWrap='CJK')
    if summary:
        draw_paragraph(c, summary[0], 18 * mm, H - 38 * mm, 116 * mm, style_hero)

    card_x, card_y = 136 * mm, H - 39 * mm
    c.setStrokeColor(colors.Color(1, 1, 1, alpha=0.45))
    c.setFillColor(colors.Color(1, 1, 1, alpha=0.10))
    c.rect(card_x, card_y, 55 * mm, 24 * mm, fill=1, stroke=1)
    c.setFillColor(white)
    c.setFont(font, 8)
    c.drawString(card_x + 5 * mm, card_y + 16 * mm, '联系方式')
    contact_y = card_y + 10 * mm
    for label, value in contact_items(data):
        c.drawString(card_x + 5 * mm, contact_y, f'{label}：{clean_visible_text(value)}')
        contact_y -= 5 * mm

    side_w = 58 * mm
    content_top = H - hero_h
    c.setFillColor(soft)
    c.rect(0, 0, side_w, content_top, fill=1, stroke=0)
    c.setStrokeColor(line)
    c.line(side_w, 0, side_w, content_top)

    style_side = ParagraphStyle('side', fontName=font, fontSize=8.2, leading=10, textColor=ink, wordWrap='CJK')
    style_body = ParagraphStyle('body', fontName=font, fontSize=8.5, leading=12, textColor=ink, wordWrap='CJK')
    style_card = ParagraphStyle('card', fontName=font, fontSize=7.5, leading=10.2, textColor=colors.HexColor('#44505a'), wordWrap='CJK')

    sidebar_x = 12 * mm
    y = content_top - 16 * mm

    def side_heading(text):
        nonlocal y
        c.setFillColor(accent)
        c.rect(sidebar_x, y - 1 * mm, 2 * mm, 5 * mm, fill=1, stroke=0)
        c.setFillColor(primary)
        c.setFont(font, 10)
        c.drawString(sidebar_x + 4 * mm, y, text)
        y -= 8 * mm

    side_heading('个人信息')
    for label, value in infer_profile(data):
        c.setFillColor(muted)
        c.setFont(font, 7.5)
        c.drawString(sidebar_x, y, clean_visible_text(label))
        h = draw_paragraph(c, value, sidebar_x, y - 4 * mm, side_w - 20 * mm, style_side)
        c.setStrokeColor(line)
        c.line(sidebar_x, y - 8.5 * mm, side_w - 8 * mm, y - 8.5 * mm)
        y -= max(11.5 * mm, h + 8 * mm)

    y -= 4 * mm
    side_heading('能力标签')
    chip_x, chip_y = sidebar_x, y
    for kw in keywords:
        tw = max(18 * mm, len(clean_visible_text(kw)) * 4.2 * mm)
        if chip_x + tw > side_w - 7 * mm:
            chip_x = sidebar_x
            chip_y -= 8 * mm
        c.setStrokeColor(accent)
        c.setFillColor(white)
        c.roundRect(chip_x, chip_y - 4 * mm, tw, 5.5 * mm, 1.2 * mm, fill=1, stroke=1)
        c.setFillColor(primary)
        c.setFont(font, 7.3)
        c.drawCentredString(chip_x + tw / 2, chip_y - 2.4 * mm, clean_visible_text(kw))
        chip_x += tw + 2 * mm

    main_x = side_w + 14 * mm
    main_w = W - main_x - 14 * mm
    y = content_top - 16 * mm

    def heading(text):
        nonlocal y
        c.setFillColor(accent)
        c.rect(main_x, y - 1 * mm, 2 * mm, 5 * mm, fill=1, stroke=0)
        c.setFillColor(primary)
        c.setFont(font, 10)
        c.drawString(main_x + 4 * mm, y, clean_visible_text(text))
        y -= 7 * mm

    if summary:
        heading('岗位摘要')
        labels = ['岗位意愿', '能力基础', '发展方向']
        card_gap = 3 * mm
        card_w = (main_w - card_gap * 2) / 3
        card_h = 25 * mm
        for idx, text in enumerate(summary[:3]):
            x = main_x + idx * (card_w + card_gap)
            c.setFillColor(white)
            c.setStrokeColor(line)
            c.rect(x, y - card_h, card_w, card_h, fill=1, stroke=1)
            c.setFillColor(accent)
            c.rect(x, y - 1.5 * mm, card_w, 1.5 * mm, fill=1, stroke=0)
            c.setFillColor(primary)
            c.setFont(font, 8.5)
            c.drawString(x + 3 * mm, y - 6 * mm, labels[idx] if idx < len(labels) else '核心优势')
            draw_paragraph(c, text, x + 3 * mm, y - 8 * mm, card_w - 6 * mm, style_card)
        y -= card_h + 10 * mm

    for section in sections:
        heading(section.get('title') or '')
        for item in section.get('items') or []:
            if y < 35 * mm:
                c.showPage()
                c.setFillColor(paper)
                c.rect(0, 0, W, H, fill=1, stroke=0)
                y = H - 18 * mm
            c.setFillColor(accent)
            c.rect(main_x + 1 * mm, y - 2.5 * mm, 2.2 * mm, 2.2 * mm, fill=1, stroke=0)
            c.setFillColor(ink)
            c.setFont(font, 9.2)
            c.drawString(main_x + 6 * mm, y, clean_visible_text(item.get('heading') or ''))
            meta = clean_visible_text(item.get('meta') or '').replace('->', '→')
            if meta:
                c.setFillColor(muted)
                c.setFont(font, 7.5)
                c.drawRightString(main_x + main_w, y, meta)
            y -= 5.5 * mm
            for bullet in item.get('bullets') or []:
                c.setFillColor(primary)
                c.circle(main_x + 7 * mm, y - 1.6 * mm, 0.75 * mm, fill=1, stroke=0)
                h = draw_paragraph(c, bullet, main_x + 11 * mm, y, main_w - 12 * mm, style_body)
                y -= h + 1.2 * mm
            y -= 2.2 * mm
        y -= 2 * mm

    if keywords and y > 24 * mm:
        heading('核心能力')
        x = main_x
        for kw in keywords:
            text = clean_visible_text(kw)
            tw = max(19 * mm, len(text) * 4.2 * mm)
            if x + tw > main_x + main_w:
                x = main_x
                y -= 8 * mm
            c.setStrokeColor(accent)
            c.setFillColor(white)
            c.roundRect(x, y - 4 * mm, tw, 5.5 * mm, 1.2 * mm, fill=1, stroke=1)
            c.setFillColor(primary)
            c.setFont(font, 7.5)
            c.drawCentredString(x + tw / 2, y - 2.4 * mm, text)
            x += tw + 2 * mm

    c.save()
    return True

def try_pdf(data, html_path, pdf_path):
    Path(pdf_path).unlink(missing_ok=True)
    ok, engine = write_pdf_chrome(html_path, pdf_path)
    if ok:
        return True, engine
    try:
        import weasyprint
        weasyprint.HTML(filename=str(html_path)).write_pdf(str(pdf_path))
        return True, 'weasyprint'
    except Exception:
        pass
    wkhtmltopdf = shutil.which('wkhtmltopdf')
    if wkhtmltopdf:
        result = subprocess.run([wkhtmltopdf, str(html_path), str(pdf_path)], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if result.returncode == 0:
            return True, 'wkhtmltopdf'
    if write_pdf_reportlab(data, pdf_path):
        return True, 'reportlab'
    return False, None


def chrome_candidates():
    candidates = [
        shutil.which('google-chrome'),
        shutil.which('chromium'),
        shutil.which('chromium-browser'),
        '/Applications/Google Chrome.app/Contents/MacOS/Google Chrome',
        '/Applications/Google Chrome Canary.app/Contents/MacOS/Google Chrome Canary',
        '/Applications/Chromium.app/Contents/MacOS/Chromium',
    ]
    return [Path(x) for x in candidates if x and Path(x).exists()]


def write_pdf_chrome(html_path, pdf_path):
    chrome = next(iter(chrome_candidates()), None)
    if not chrome:
        return False, None
    Path(pdf_path).unlink(missing_ok=True)
    profile_dir = Path('/tmp') / f'resume-chrome-profile-{os.getpid()}'
    if profile_dir.exists():
        shutil.rmtree(profile_dir, ignore_errors=True)
    profile_dir.mkdir(parents=True, exist_ok=True)
    cmd = [
        str(chrome),
        '--headless=new',
        '--disable-gpu',
        '--disable-background-networking',
        '--disable-component-update',
        '--disable-default-apps',
        '--disable-sync',
        '--hide-scrollbars',
        '--no-first-run',
        '--no-default-browser-check',
        '--print-to-pdf-no-header',
        f'--user-data-dir={profile_dir}',
        f'--print-to-pdf={pdf_path}',
        Path(html_path).resolve().as_uri(),
    ]
    proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    try:
        proc.communicate(timeout=25)
    except subprocess.TimeoutExpired:
        proc.terminate()
        try:
            proc.communicate(timeout=3)
        except subprocess.TimeoutExpired:
            proc.kill()
            proc.communicate(timeout=3)
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)
    if Path(pdf_path).exists() and Path(pdf_path).stat().st_size > 10000:
        return True, 'chrome'
    return False, None


def main():
    parser = argparse.ArgumentParser(description='Render Chinese graduate resume HTML/PDF artifacts from JSON.')
    parser.add_argument('input_json')
    parser.add_argument('--out-dir', default='.')
    parser.add_argument('--formats', default='html')
    args = parser.parse_args()

    data = json.loads(Path(args.input_json).read_text(encoding='utf-8'))
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    basics = data.get('basics', {})
    target = data.get('target', {})
    stem = safe_name('-'.join(x for x in [basics.get('name'), target.get('role') or basics.get('title'), '简历'] if x))
    formats = {x.strip().lower() for x in args.formats.split(',') if x.strip()}
    if 'docx' in formats:
        print('WARNING: 本 skill 已取消 Word/DOCX 交付，已忽略 docx；只生成 HTML 和 PDF。', file=sys.stderr)
        formats.discard('docx')

    outputs = []
    html_path = out_dir / f'{stem}.html'
    if formats & {'html', 'pdf'}:
        html_path.write_text(render_html(data), encoding='utf-8')
        outputs.append(html_path)
    if 'pdf' in formats:
        pdf_path = out_dir / f'{stem}.pdf'
        ok, _ = try_pdf(data, html_path, pdf_path)
        if ok:
            outputs.append(pdf_path)
        else:
            print('WARNING: 本机未检测到可用 PDF 生成链路，未生成 PDF。已保留 HTML；请用浏览器打印为 PDF。', file=sys.stderr)
    for output in outputs:
        print(output)


if __name__ == '__main__':
    main()
